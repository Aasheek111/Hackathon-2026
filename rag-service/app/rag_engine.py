"""PDF ingestion, vector storage and tutorial generation.

The NeuroLearn RAG service. Everything that touches the LLM or the vector store
lives here, so the route layer stays thin and testable.

The piece that matters for this platform: generation takes a **learning mode**
(TEXT / AUDIO / VISUAL / AR - the same enum the Prisma schema uses). When the CV
service detects disengagement and the platform switches mode, the same unit can
be regenerated in the new mode from the same source PDF.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import struct
import time
from pathlib import Path

import requests
from dotenv import load_dotenv
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

# override=True: docker-compose sets GROQ_API_KEY="" in the container when
# it isn't exported in the host shell (the ${GROQ_API_KEY:-} fallback), and
# python-dotenv otherwise leaves an already-set env var alone - which would
# silently shadow a real key placed in this file, the documented way to
# configure it.
load_dotenv(override=True)

# --- paths ------------------------------------------------------------------
# Resolved from this file rather than the working directory, so the server
# behaves the same whether it is started from the project root or elsewhere.
BASE_DIR = Path(__file__).resolve().parent.parent
PDF_DIR = BASE_DIR / "uploads" / "pdfs"
IMAGE_DIR = BASE_DIR / "uploads" / "images"
AUDIO_DIR = BASE_DIR / "uploads" / "audio"
VECTOR_DIR = BASE_DIR / "vector_store"
STATIC_DIR = BASE_DIR / "static"

CHUNK_SIZE = 500
CHUNK_OVERLAP = 50
TOP_K = 4

DEFAULT_QUERY = "Explain the key concepts of this unit simply"

SYSTEM_PROMPT = (
    "You are a special education tutor. You receive textbook chunks and a student diagnosis.\n"
    "Generate a response in JSON format with keys: tutorial_text, audio_script, "
    "visual_suggestion, steps, quiz, teacher_note.\n"
    "'steps' is a list of 2-4 objects, each with 'concept' (a short heading), 'explanation' "
    "(1-2 plain-language sentences), and 'example' (one concrete example) - break the unit's "
    "content into that many bite-sized lessons, each grounded in the textbook chunks provided.\n"
    "'quiz' should have 2-3 questions, each with 4 options, grounded in different parts of the "
    "textbook chunks - not all testing the same fact.\n"
    "Before finalizing, re-read every quiz question against its own options and the textbook "
    "chunks: the 'correct' value must be the option that actually and specifically answers that "
    "exact question - not a fact from a different question, and not merely a true statement "
    "from the text. Double-check each one before including it.\n"
    "If the diagnosis says 'Struggling', use grade-1 vocabulary and only 2 MCQ options per question. "
    "Keep tutorial_text under 100 words."
)

# Matches the LearningMode enum in backend/prisma/schema.prisma. The platform
# already decides the mode from engagement; this is how that decision reaches the
# content.
MODE_GUIDANCE = {
    "TEXT": (
        "The learner is in TEXT mode. Write tutorial_text as short bullet points in plain "
        "language. Keep audio_script brief."
    ),
    "AUDIO": (
        "The learner is in AUDIO mode and may not be looking at the screen. Put the real "
        "teaching in audio_script: full narration, describing anything visual out loud, in "
        "short spoken sentences. Keep tutorial_text to a summary."
    ),
    "VISUAL": (
        "The learner is in VISUAL mode. Lead with what to draw: visual_suggestion must "
        "describe one concrete, simple diagram or chart a teacher could sketch in 30 seconds. "
        "Keep the wording minimal and concrete."
    ),
    "AR": (
        "The learner is in AR mode and has disengaged from reading. Reduce to the single most "
        "important idea, make the quiz playful and physical, and describe a visual that could "
        "become a 3D object."
    ),
}
VALID_MODES = tuple(MODE_GUIDANCE)


def ensure_directories() -> None:
    """Called on startup so a fresh clone works without any manual mkdir."""
    for directory in (PDF_DIR, IMAGE_DIR, AUDIO_DIR, VECTOR_DIR, STATIC_DIR):
        directory.mkdir(parents=True, exist_ok=True)


def index_path(unit_id: int) -> Path:
    return VECTOR_DIR / f"index_unit_{unit_id}"


def pdf_path(unit_id: int) -> Path:
    return PDF_DIR / f"unit_{unit_id}.pdf"


def unit_is_processed(unit_id: int) -> bool:
    """A FAISS index is a directory holding index.faiss and index.pkl."""
    return (index_path(unit_id) / "index.faiss").exists()


def processed_units() -> list[int]:
    units = []
    for entry in VECTOR_DIR.glob("index_unit_*"):
        match = re.fullmatch(r"index_unit_(\d+)", entry.name)
        if match and (entry / "index.faiss").exists():
            units.append(int(match.group(1)))
    return sorted(units)


# the value shipped in .env.example - treating it as "configured" would mean the
# app looks healthy right up until the first request fails
PLACEHOLDER_KEYS = {"", "your-key-here", "changeme"}


# --- embeddings ---------------------------------------------------------------
# Runs LOCALLY (fastembed / ONNX, ~66MB model cached in the image). Deliberately
# not a hosted embedding API: Gemini's embedding endpoint shares the same
# free-tier request budget that kept exhausting mid-generation, and embeddings
# are called once per chunk, so a 40-chunk PDF could burn a day's quota on
# indexing alone. Local means no key, no quota, no network, and it keeps
# working when every external provider is rate-limited.
EMBEDDING_MODEL = "BAAI/bge-small-en-v1.5"
_embeddings_cache = None


def get_embeddings():
    """Cached because constructing it loads the ONNX model from disk."""
    global _embeddings_cache
    if _embeddings_cache is None:
        from langchain_community.embeddings import FastEmbedEmbeddings

        _embeddings_cache = FastEmbedEmbeddings(model_name=EMBEDDING_MODEL)
    return _embeddings_cache


# --- Groq (primary text generation) ------------------------------------------
#
# Verified live (real key, curl against the actual API) before wiring this
# in: llama-3.3-70b-versatile + response_format={"type": "json_object"}
# (OpenAI-compatible, confirmed working). Groq's free tier is dramatically
# more generous than Gemini's for text generation - 1,000 requests/day and
# 30/minute, versus Gemini's observed 20/day - which is the actual bottleneck
# this pipeline has hit repeatedly. Groq is now the ONLY text engine - the
# Gemini fallback was removed because falling back to a 20-request/day quota
# just converted a transient Groq blip into a confusing Gemini error. Visuals
# come from Unsplash and embeddings run locally, so Groq is the only AI
# provider this service talks to.

GROQ_CHAT_MODEL = "llama-3.3-70b-versatile"
GROQ_CHAT_URL = "https://api.groq.com/openai/v1/chat/completions"

# Observed live: a transient Docker-network hiccup (ConnectionRefusedError
# reaching api.groq.com) made a single call silently fall through to Gemini,
# which then hit its already-exhausted daily quota and failed the whole job -
# even though the very next attempt to Groq, seconds later, succeeded. A
# couple of retries turns that into a few seconds' delay instead of an entire
# job failing over a blip that had nothing to do with either provider's quota.
GROQ_REQUEST_ATTEMPTS = 3
GROQ_RETRY_DELAY_SECONDS = 2


def groq_key_present() -> bool:
    return os.getenv("GROQ_API_KEY", "").strip().strip('"') not in PLACEHOLDER_KEYS


def any_llm_configured() -> bool:
    """Gate for the offline-extraction fallbacks. Groq is the only text
    provider now, so this is simply "is Groq configured".
    """
    return groq_key_present()


def _invoke_groq_json(system_prompt: str, user_prompt: str) -> str | None:
    if not groq_key_present():
        return None
    for attempt in range(GROQ_REQUEST_ATTEMPTS):
        try:
            response = requests.post(
                GROQ_CHAT_URL,
                headers={"Authorization": f"Bearer {os.getenv('GROQ_API_KEY')}", "Content-Type": "application/json"},
                json={
                    "model": GROQ_CHAT_MODEL,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    "response_format": {"type": "json_object"},
                    "temperature": 0.3,
                },
                timeout=60,
            )
            # 429 is Groq's per-minute rate limit, not a real failure - it
            # clears within seconds, so it's worth the same retry as a
            # network blip. Any other error status (bad request, auth,
            # safety block) needs a code/account fix, not a retry.
            if response.status_code == 429 and attempt < GROQ_REQUEST_ATTEMPTS - 1:
                time.sleep(GROQ_RETRY_DELAY_SECONDS)
                continue
            response.raise_for_status()
            return response.json()["choices"][0]["message"]["content"]
        except (requests.exceptions.ConnectionError, requests.exceptions.Timeout):
            # transient network blip (observed live: ConnectionRefusedError
            # reaching api.groq.com that succeeded again seconds later) -
            # worth retrying rather than immediately burning Gemini's much
            # tighter quota as a fallback
            if attempt < GROQ_REQUEST_ATTEMPTS - 1:
                time.sleep(GROQ_RETRY_DELAY_SECONDS)
                continue
            return None
        except Exception:
            return None
    return None


def invoke_json(system_prompt: str, user_prompt: str) -> str:
    """The one place every JSON-generation call in this module goes through.

    Groq ONLY. The Gemini fallback was removed deliberately: Gemini's free tier
    allows 20 generate-content requests per DAY, so once Groq was briefly
    unavailable the pipeline would silently drop onto Gemini, exhaust it within
    a single unit, and then surface a confusing "gemini quota exceeded" error
    for a failure that had nothing to do with Gemini. One provider means the
    error you see is the error that actually happened.
    """
    result = _invoke_groq_json(system_prompt, user_prompt)
    if result is not None:
        return result
    if not groq_key_present():
        raise RuntimeError("GROQ_API_KEY is not set. Copy .env.example to .env and add your key.")
    raise RuntimeError(
        "Groq did not return a result (rate limit or network). Free tier allows "
        "1,000 requests and 100,000 tokens per day - if a large document exhausts "
        "the daily token budget, generation resumes on the next reset."
    )


# --- ingestion ---------------------------------------------------------------


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def extract_text(path: Path) -> str:
    """Pull the text layer out of a teacher's document.

    Handles PDF, Word (.docx) and plain text/markdown. A scanned PDF or a
    photo has no text layer, so this legitimately returns almost nothing -
    the caller checks and says so rather than silently indexing an empty
    document. Images (.png/.jpg) are NOT supported: they would need OCR,
    which is a genuinely different pipeline, so we reject them at upload with
    a clear message instead of pretending to read them.
    """
    suffix = path.suffix.lower()

    if suffix == ".docx":
        from docx import Document  # python-docx

        document = Document(str(path))
        parts = [p.text for p in document.paragraphs]
        # Tables carry real syllabus content (unit breakdowns, mark schemes),
        # so they're pulled in too rather than silently dropped.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(p for p in parts if p.strip()).strip()

    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore").strip()

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def split_text(text: str) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        # try to break on paragraphs first, then sentences, then words, so a
        # chunk is usually a whole idea rather than half a sentence
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_text(text)


def build_index(unit_id: int, chunks: list[str]) -> int:
    """Embed the chunks and persist a FAISS index for this unit.

    Without an API key there is nothing to call OpenAI's embeddings endpoint
    with, so a placeholder local embedding is used instead. Those vectors are
    never used for offline retrieval (`retrieve` reads the docstore directly
    in that case) - they only exist so FAISS has something to store the real
    chunk text against. This keeps ingestion working offline, symmetric with
    `generate_tutorial`'s existing offline fallback.
    """
    metadatas = [{"unit_id": unit_id, "chunk": position} for position in range(len(chunks))]
    embeddings = get_embeddings()
    store = FAISS.from_texts(chunks, embeddings, metadatas=metadatas)
    target = index_path(unit_id)
    target.mkdir(parents=True, exist_ok=True)
    store.save_local(str(target))
    return len(chunks)


def load_index(unit_id: int) -> FAISS:
    # allow_dangerous_deserialization is required because FAISS stores its
    # docstore as a pickle. Safe here: we only ever load indexes this app wrote.
    return FAISS.load_local(
        str(index_path(unit_id)),
        get_embeddings(),
        allow_dangerous_deserialization=True,
    )


def retrieve(unit_id: int, query: str, k: int = TOP_K) -> list[str]:
    """Top-k chunks for this query.

    Embedding runs locally now, so this no longer depends on any API key or
    quota. Falls back to raw docstore chunks only if the index can't be read.
    """
    try:
        store = load_index(unit_id)
    except Exception:
        return raw_chunks(unit_id, k)
    return [document.page_content for document in store.similarity_search(query, k=k)]


def raw_chunks(unit_id: int, limit: int = TOP_K) -> list[str]:
    """Read chunks back without embedding anything (offline path)."""
    import pickle

    store_file = index_path(unit_id) / "index.pkl"
    if not store_file.exists():
        return []
    try:
        docstore, _ = pickle.loads(store_file.read_bytes())
        documents = list(getattr(docstore, "_dict", {}).values())
        return [document.page_content for document in documents[:limit]]
    except Exception:
        return []


# --- generation ---------------------------------------------------------------


def _coerce_json(raw: str) -> dict:
    """Parse the model's reply, tolerating a fenced code block."""
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    return json.loads(text)


def _normalise(payload: dict) -> dict:
    """Guarantee the exact response shape, whatever the model returned.

    The frontend and any marking script depend on these keys existing, so a
    missing one becomes an empty value rather than a KeyError downstream.
    """
    quiz = []
    for item in payload.get("quiz") or []:
        if not isinstance(item, dict):
            continue
        options = [str(option) for option in (item.get("options") or [])]
        correct = str(item.get("correct", options[0] if options else ""))
        # A worse failure than a wrong-but-plausible answer: the model naming
        # a "correct" value that isn't even one of its own options, which
        # would make the question unanswerable correctly. Drop it rather than
        # ship a quiz item that can never be marked right.
        if options and correct not in options:
            continue
        quiz.append({"question": str(item.get("question", "")), "options": options, "correct": correct})

    steps = []
    for item in payload.get("steps") or []:
        if not isinstance(item, dict):
            continue
        steps.append(
            {
                "concept": str(item.get("concept", "")),
                "explanation": str(item.get("explanation", "")),
                "example": str(item.get("example", "")),
            }
        )

    return {
        "tutorial_text": str(payload.get("tutorial_text", "")),
        "audio_script": str(payload.get("audio_script", "")),
        "visual_suggestion": str(payload.get("visual_suggestion", "")),
        "steps": steps,
        "quiz": quiz,
        "teacher_note": str(payload.get("teacher_note", "")),
    }


def _first_sentences(text: str, count: int = 3) -> str:
    parts = re.split(r"(?<=[.!?])\s+", text.strip())
    return " ".join(parts[:count]).strip()


def offline_tutorial(unit_id: int, chunks: list[str], learning_mode: str) -> dict:
    """A usable tutorial with no API key, built from the retrieved text itself.

    Not as good as the model - it summarises by extraction rather than by
    understanding - but it keeps a demo alive when the key is missing or the
    venue's wifi is down, and it is clearly labelled so nobody mistakes it for
    the real thing.
    """
    body = " ".join(chunks) if chunks else ""
    summary = _first_sentences(body, 3) or "No text was found for this unit."
    headline = _first_sentences(body, 1) or summary

    usable_chunks = [c for c in chunks if c.strip()] or ([body] if body else [])
    steps = []
    for i, chunk in enumerate(usable_chunks[:4]):
        sentences = [s for s in re.split(r"(?<=[.!?])\s+", chunk.strip()) if s]
        steps.append(
            {
                "concept": f"Part {i + 1}",
                "explanation": " ".join(sentences[:2]) or chunk[:200],
                "example": sentences[2] if len(sentences) > 2 else "",
            }
        )
    if not steps:
        steps = [{"concept": "Overview", "explanation": summary, "example": ""}]

    # Build each question from a different chunk when there are enough of
    # them, using OTHER chunks' headlines as distractors - every option is
    # real extracted text, nothing invented.
    headlines = [
        (_first_sentences(c, 1) or c)[:80] for c in usable_chunks if c.strip()
    ] or [headline[:80]]
    quiz = []
    for i, correct in enumerate(headlines[:3]):
        distractors = [h for j, h in enumerate(headlines) if j != i][:3]
        options = [correct] + distractors
        if len(options) < 2:
            options.append("Something not covered in this unit")
        quiz.append(
            {
                "question": f"Which of these is discussed in this unit (part {i + 1})?",
                "options": options,
                "correct": correct,
            }
        )

    return {
        "tutorial_text": f"Key points from this unit:\n• {summary}",
        "audio_script": (
            f"Here is what this unit covers. {summary} "
            "Take your time, and ask if you would like it again."
        ),
        "visual_suggestion": (
            "A simple labelled diagram of the main idea above - one box per key term, "
            "connected left to right."
        ),
        "steps": steps,
        "quiz": quiz,
        "teacher_note": (
            "Offline mode: built by extracting sentences, not by the model. "
            "Set GROQ_API_KEY for an adapted tutorial."
        ),
        "source_chunks": len(chunks),
        "learning_mode": learning_mode,
        "offline": True,
    }


def grade_prefix(grade_level: str | None) -> str:
    """A one-line instruction pinning the target education level, prepended to
    every generation system prompt. Admin-set (see backend AppConfig) so the
    whole platform can be retargeted (Nursery -> Grade 5, etc.) without code
    changes. Empty when no level is configured, leaving prompts unchanged.
    """
    grade = (grade_level or "").strip()
    if not grade:
        return ""
    return (
        f"TARGET LEARNERS: {grade}-level students. Use vocabulary, sentence length, examples, "
        f"and question difficulty appropriate for {grade}. Never exceed that level.\n\n"
    )


def generate_tutorial(
    unit_id: int, student_diagnosis: str | None, learning_mode: str = "TEXT", grade_level: str | None = None
) -> dict:
    """Retrieve the relevant chunks and turn them into an adapted tutorial."""
    mode = (learning_mode or "TEXT").upper()
    if mode not in VALID_MODES:
        mode = "TEXT"

    query = (student_diagnosis or "").strip() or DEFAULT_QUERY
    chunks = retrieve(unit_id, query)

    # no key: fall back rather than fail. A dead demo helps nobody.
    if not any_llm_configured():
        return offline_tutorial(unit_id, chunks, mode)

    context = "\n\n---\n\n".join(chunks) if chunks else "(no textbook content found)"
    user_prompt = (
        f"Student diagnosis: {student_diagnosis or 'not provided'}\n"
        f"Learning mode: {mode}. {MODE_GUIDANCE[mode]}\n\n"
        f"Textbook chunks:\n{context}\n\n"
        "Return only the JSON object, with no commentary."
    )

    raw = invoke_json(grade_prefix(grade_level) + SYSTEM_PROMPT, user_prompt)

    result = _normalise(_coerce_json(raw))
    # useful for the teacher and for debugging retrieval quality
    result["source_chunks"] = len(chunks)
    result["learning_mode"] = mode
    result["offline"] = False
    return result


# --- image generation --------------------------------------------------------


# Kept for the mimetype of images written to disk by earlier versions, which
# are still served from /static/images. Nothing new is written there - Unsplash
# images are hotlinked.
_MIME_EXTENSIONS = {"image/png": ".png", "image/jpeg": ".jpg", "image/webp": ".webp"}


# pollinations.ai was removed here deliberately. It generated each image on
# demand, which meant waiting on image synthesis per lesson, and it throttled
# hard under a curriculum-sized burst (observed live: 1 of 21 succeeded, the
# rest 429/timeout, needing backoff retries that made a full unit take
# minutes). Unsplash returns a real photo in a single fast search call and is
# hotlinked, so nothing is generated, downloaded, or stored.


UNSPLASH_SEARCH_URL = "https://api.unsplash.com/search/photos"

# Unsplash is a photo SEARCH API, not an image generator - so the AI's
# "visual_suggestion" (written as a prompt for an illustrator) has to be
# reduced to a few keywords or it matches nothing. Strip the styling
# boilerplate and keep the first handful of content words.
_UNSPLASH_STOPWORDS = {
    "a", "an", "the", "of", "for", "to", "and", "or", "with", "in", "on", "that", "this",
    "showing", "show", "shows", "illustration", "illustrating", "diagram", "picture",
    "image", "drawing", "draw", "simple", "clean", "educational", "student", "students",
    "depicting", "depict", "visual", "chart", "graphic", "showing",
}
UNSPLASH_MAX_KEYWORDS = 5


def unsplash_key_present() -> bool:
    return os.getenv("UNSPLASH_ACCESS_KEY", "").strip().strip('"') not in PLACEHOLDER_KEYS


def _unsplash_keywords(prompt: str) -> list[str]:
    words = re.findall(r"[A-Za-z][A-Za-z'-]+", prompt.lower())
    return [w for w in words if w not in _UNSPLASH_STOPWORDS and len(w) > 2]


def _unsplash_queries(prompt: str) -> list[str]:
    """Progressively BROADER search terms, most specific first.

    A photo library has nothing for "barriers faced by nepalese entrepreneurs",
    which is exactly why lessons on abstract topics were coming back with no
    picture at all. Narrowing 5 keywords -> 2 -> 1 -> a generic topic term means
    a specific match wins when one exists, and every other lesson still gets a
    relevant image instead of an empty placeholder.
    """
    kw = _unsplash_keywords(prompt)
    queries = []
    for n in (UNSPLASH_MAX_KEYWORDS, 3, 2, 1):
        if len(kw) >= n:
            q = " ".join(kw[:n])
            if q not in queries:
                queries.append(q)
    if not queries:
        queries.append((prompt.strip()[:60] or "education"))
    queries.append("education learning classroom")  # never leave a lesson blank
    return queries


def _unsplash_search(query: str) -> str | None:
    try:
        response = requests.get(
            UNSPLASH_SEARCH_URL,
            params={"query": query, "per_page": 1, "orientation": "landscape"},
            headers={
                "Authorization": f"Client-ID {os.getenv('UNSPLASH_ACCESS_KEY', '').strip().strip(chr(34))}",
                "Accept-Version": "v1",
            },
            timeout=15,
        )
        if response.status_code != 200:
            return None
        results = response.json().get("results") or []
        if not results:
            return None
        urls = results[0].get("urls") or {}
        return urls.get("regular") or urls.get("small") or urls.get("full")
    except Exception:
        return None


def _fetch_via_unsplash(prompt: str) -> str | None:
    """Returns a HOTLINKED Unsplash CDN URL (not downloaded bytes).

    Unsplash's API guidelines require hotlinking their CDN rather than
    re-hosting copies, and it's also far faster than image synthesis - a real
    photo arrives in one search call and we store nothing.
    """
    if not unsplash_key_present():
        return None
    for query in _unsplash_queries(prompt):
        hit = _unsplash_search(query)
        if hit:
            return hit
    return None


def generate_visual_image(prompt: str, unit_id: int) -> str | None:
    """Turn a visual_suggestion (a hint written for a human artist) into an
    actual picture.

    Unsplash only. Both earlier generators are gone: pollinations was slow and
    throttled under a curriculum-sized burst, and Gemini's image model needs
    billing enabled so it never actually produced anything on the free tier -
    it just added a wasted round-trip per lesson. Progressive keyword
    broadening (see _unsplash_queries) means every lesson resolves to a photo,
    which is what fixes lessons rendering "No picture for this lesson yet".

    Returns an absolute Unsplash CDN URL; the frontend's resolveMediaUrl()
    also still handles the relative /static/images/... paths of anything
    generated before this change.

    Never raises - if it comes back empty the caller just shows text, same
    contract as `offline_tutorial`. A dead demo helps nobody.
    """
    if not (prompt or "").strip():
        return None
    return _fetch_via_unsplash(prompt)


# --- full-document curriculum generation --------------------------------------
#
# generate_tutorial() above deliberately only sees the top-K similarity-matched
# chunks (fine for "answer this student's specific gap"). A curriculum needs
# the OPPOSITE: every chunk, in order, actually represented in the lesson
# plan - the exact gap TODO.md Phase 2 exists to close.

CURRICULUM_PLAN_PROMPT = (
    "You are a curriculum designer. Below is a full textbook document, broken into "
    "numbered chunks. Design a complete lesson plan that covers EVERY educationally "
    "meaningful part of this document - do not skip sections, do not stop after the "
    "introduction, and do not artificially limit yourself to a small fixed number of "
    "lessons. The number of lessons must reflect how much distinct content is "
    "actually in the document: a short document might need only 2-3 lessons, a long "
    "one might need 15 or more.\n\n"
    "Group the chunks into an ordered list of lessons. Each lesson must be a "
    "self-contained, teachable concept and must reference the exact chunk range "
    "(start and end chunk numbers, inclusive) it is grounded in. Chunk ranges must "
    "not overlap, must stay in ascending order, and together should cover as much of "
    "the document as is educationally meaningful (it is fine to skip a chunk that is "
    "clearly a table of contents, references list, or similar non-content section).\n\n"
    "Return JSON: {\"title\": <overall subject/topic title>, "
    "\"lessons\": [{\"title\": <short lesson heading>, \"chunk_start\": <int>, "
    "\"chunk_end\": <int>}]}. Return only the JSON object, no commentary."
)

LESSON_SYSTEM_PROMPT = (
    "You are a special education tutor writing ONE lesson of a larger curriculum, "
    "grounded only in the textbook excerpt provided - do not introduce facts that "
    "are not in the excerpt.\n"
    "Generate JSON with keys: explanation, example, needs_visual, visual_suggestion, "
    "knowledge_check.\n"
    "'explanation': 2-4 plain-language sentences teaching the concept.\n"
    "'example': one concrete example grounded in the excerpt (empty string if none fits).\n"
    "'needs_visual': boolean - true only when a picture would genuinely help understand "
    "THIS specific lesson. Do not default to true for every lesson - most lessons about "
    "abstract or purely verbal content do not need one.\n"
    "'visual_suggestion': if needs_visual is true, describe one concrete, simple diagram "
    "or illustration; empty string otherwise.\n"
    "'knowledge_check': {\"question\", \"options\" (2-4 strings), \"correct\"} - a single "
    "check-for-understanding question. Before finalizing, verify 'correct' is exactly one "
    "of the 'options' values."
)

FINAL_ASSESSMENT_SYSTEM_PROMPT = (
    "You write a final assessment for a completed learning curriculum, covering its "
    "lessons. Write up to 10 multiple-choice questions - fewer if the curriculum is "
    "short, never pad with repetitive questions testing the same fact. Each question "
    "needs exactly 4 options and a 'correct' value that is exactly one of them, "
    "grounded in a different lesson where possible so the assessment covers the whole "
    "curriculum rather than one section repeatedly.\n"
    "Return JSON: {\"questions\": [{\"question\", \"options\", \"correct\"}]}. Only the JSON."
)


def all_chunks(unit_id: int) -> list[dict]:
    """Every chunk this unit's PDF was split into, in original document order,
    each tagged with its position.

    `retrieve()` intentionally only returns the top-k similarity matches for a
    single student query - wrong input for a curriculum planner, which needs
    to see the whole document to guarantee full coverage.
    """
    import pickle

    store_file = index_path(unit_id) / "index.pkl"
    if not store_file.exists():
        return []
    try:
        docstore, _ = pickle.loads(store_file.read_bytes())
        documents = list(getattr(docstore, "_dict", {}).values())
    except Exception:
        return []
    ordered = sorted(documents, key=lambda d: d.metadata.get("chunk", 0))
    return [
        {"position": d.metadata.get("chunk", i), "text": d.page_content}
        for i, d in enumerate(ordered)
    ]


def plan_curriculum(unit_id: int, grade_level: str | None = None) -> dict:
    """One LLM call over the WHOLE document, producing an ordered lesson plan
    whose lesson count is derived from actual content (TODO.md Phase 2 - never
    a hardcoded number).
    """
    chunks = all_chunks(unit_id)
    if not chunks:
        raise ValueError("No content found for this unit")

    if not any_llm_configured():
        # Offline fallback: one lesson per chunk, so the document is still
        # fully represented even without a model to plan groupings.
        return {
            "title": "Curriculum",
            "lessons": [
                {"title": f"Part {c['position'] + 1}", "chunk_start": c["position"], "chunk_end": c["position"]}
                for c in chunks
            ],
            "total_chunks": len(chunks),
        }

    numbered = "\n\n".join(f"[chunk {c['position']}] {c['text']}" for c in chunks)
    plan = _coerce_json(invoke_json(grade_prefix(grade_level) + CURRICULUM_PLAN_PROMPT, numbered))

    lessons = []
    for item in plan.get("lessons") or []:
        try:
            start = int(item["chunk_start"])
            end = int(item["chunk_end"])
        except (KeyError, TypeError, ValueError):
            continue
        lessons.append({"title": str(item.get("title", "")) or f"Chunks {start}-{end}", "chunk_start": start, "chunk_end": end})

    if not lessons:
        # The model returned nothing usable - fall back to one lesson per
        # chunk rather than silently producing an empty curriculum.
        lessons = [
            {"title": f"Part {c['position'] + 1}", "chunk_start": c["position"], "chunk_end": c["position"]}
            for c in chunks
        ]
    else:
        lessons = _fill_coverage_gaps(lessons, chunks)

    return {"title": str(plan.get("title") or "Curriculum"), "lessons": lessons, "total_chunks": len(chunks)}


# How many consecutive uncovered chunks to bundle into one "gap" lesson - keeps
# a big missed tail from exploding into dozens of tiny one-chunk lessons.
COVERAGE_GAP_LESSON_SIZE = 4


def _fill_coverage_gaps(lessons: list[dict], chunks: list[dict]) -> list[dict]:
    """Guarantee the WHOLE document is represented (TODO.md Phase 2 / user
    ask: "cover all, not a portion of the PDF").

    The planner LLM sometimes stops early - e.g. it plans lessons for chunks
    0-12 of a 40-chunk document and silently drops the rest, or leaves an
    interior gap. This scans for chunk positions no lesson covers and appends
    grouped "continued" lessons for them. Pure bookkeeping over chunk ranges,
    no extra LLM call here (each appended lesson's text is still generated
    later, grounded in its own chunks, exactly like a planned one).
    """
    positions = sorted(c["position"] for c in chunks)
    covered: set[int] = set()
    for lesson in lessons:
        for p in range(lesson["chunk_start"], lesson["chunk_end"] + 1):
            covered.add(p)

    missing = [p for p in positions if p not in covered]
    if not missing:
        return lessons

    # Group consecutive missing positions into a few lessons rather than one
    # per chunk.
    extra: list[dict] = []
    run_start = missing[0]
    prev = missing[0]
    part = len(lessons) + 1

    def flush(start: int, end: int, part_no: int) -> None:
        # Split a long run into COVERAGE_GAP_LESSON_SIZE-chunk lessons.
        s = start
        n = part_no
        while s <= end:
            e = min(s + COVERAGE_GAP_LESSON_SIZE - 1, end)
            extra.append({"title": f"Part {n} (continued)", "chunk_start": s, "chunk_end": e})
            s = e + 1
            n += 1

    for p in missing[1:]:
        if p == prev + 1:
            prev = p
            continue
        flush(run_start, prev, part)
        part += 1
        run_start = p
        prev = p
    flush(run_start, prev, part)

    # Keep the reading order sensible: sort all lessons by where they start.
    return sorted(lessons + extra, key=lambda l: l["chunk_start"])


def generate_lesson_content(
    unit_id: int, lesson_title: str, chunk_start: int, chunk_end: int, grade_level: str | None = None
) -> dict:
    """Content for ONE lesson, grounded only in its own chunk range - keeps
    each generation call small and traceable instead of re-processing the
    whole document per lesson.
    """
    chunks = all_chunks(unit_id)
    scoped = [c["text"] for c in chunks if chunk_start <= c["position"] <= chunk_end]
    context = "\n\n".join(scoped) or "(no content found for this range)"

    if not any_llm_configured():
        sentences = [s for s in re.split(r"(?<=[.!?])\s+", context.strip()) if s]
        return {
            "explanation": " ".join(sentences[:3]) or context[:300],
            "example": sentences[3] if len(sentences) > 3 else "",
            "needs_visual": False,
            "visual_suggestion": "",
            "knowledge_check": None,
        }

    user_prompt = f"Lesson title: {lesson_title}\n\nTextbook excerpt:\n{context}\n\nReturn only the JSON object."
    data = _coerce_json(invoke_json(grade_prefix(grade_level) + LESSON_SYSTEM_PROMPT, user_prompt))

    knowledge_check = None
    kc = data.get("knowledge_check")
    if isinstance(kc, dict):
        options = [str(o) for o in (kc.get("options") or [])]
        correct = str(kc.get("correct", options[0] if options else ""))
        if options and correct in options:
            knowledge_check = {"question": str(kc.get("question", "")), "options": options, "correct": correct}

    return {
        "explanation": str(data.get("explanation", "")),
        "example": str(data.get("example", "")),
        "needs_visual": bool(data.get("needs_visual")),
        "visual_suggestion": str(data.get("visual_suggestion", "")),
        "knowledge_check": knowledge_check,
    }


def generate_final_assessment(lesson_titles: list[str], grade_level: str | None = None) -> list[dict]:
    """Up to 10 MCQs covering the whole curriculum. Returns [] rather than
    raising when there's no API key - a curriculum without a final assessment
    still has real value; the frontend just won't show that step.
    """
    if not any_llm_configured() or not lesson_titles:
        return []

    context = "\n".join(f"- {title}" for title in lesson_titles)
    data = _coerce_json(
        invoke_json(
            grade_prefix(grade_level) + FINAL_ASSESSMENT_SYSTEM_PROMPT,
            f"Curriculum lessons:\n{context}\n\nReturn only the JSON object.",
        )
    )

    questions = []
    for item in (data.get("questions") or [])[:10]:
        options = [str(o) for o in (item.get("options") or [])]
        correct = str(item.get("correct", options[0] if options else ""))
        if options and correct in options:
            questions.append({"question": str(item.get("question", "")), "options": options, "correct": correct})
    return questions


# --- text-to-speech -----------------------------------------------------------
#
# Verified live during implementation (Docker exec against the real API):
# gemini-2.5-flash-preview-tts + the "Achernar" prebuilt voice both work and
# return real audio. The originally-requested model name
# ("gemini-3.1-flash-tts-preview") is not a real model as of this session's
# knowledge - this is the confirmed-working substitute, not an assumption.
# The API returns raw PCM (audio/L16, 24kHz, mono, 16-bit) rather than a
# browser-playable container, so _wrap_pcm_as_wav() adds a WAV header before
# saving.
#
# Groq also offers TTS (canopylabs/orpheus-v1-english, verified live via curl
# with a real key) and is tried FIRST - it already returns a proper WAV, no
# post-processing needed. As of this session, that model returns
# `model_terms_required` until the org admin accepts its terms in the Groq
# console (https://console.groq.com/playground?model=canopylabs%2Forpheus-v1-english)
# - this is not a code bug, and the function degrades to Gemini automatically
# in the meantime.

GEMINI_TTS_MODEL = "gemini-2.5-flash-preview-tts"
DEFAULT_TTS_VOICE = "Achernar"
TTS_SAMPLE_RATE = 24000

GROQ_TTS_MODEL = "canopylabs/orpheus-v1-english"
GROQ_TTS_VOICE = "troy"
GROQ_TTS_URL = "https://api.groq.com/openai/v1/audio/speech"


def _wrap_pcm_as_wav(pcm_data: bytes, sample_rate: int = TTS_SAMPLE_RATE, channels: int = 1, bits_per_sample: int = 16) -> bytes:
    byte_rate = sample_rate * channels * bits_per_sample // 8
    block_align = channels * bits_per_sample // 8
    header = struct.pack(
        "<4sI4s4sIHHIIHH4sI",
        b"RIFF", 36 + len(pcm_data), b"WAVE",
        b"fmt ", 16, 1, channels, sample_rate, byte_rate, block_align, bits_per_sample,
        b"data", len(pcm_data),
    )
    return header + pcm_data


def _generate_speech_via_groq(text: str) -> bytes | None:
    if not groq_key_present():
        return None
    for attempt in range(GROQ_REQUEST_ATTEMPTS):
        try:
            response = requests.post(
                GROQ_TTS_URL,
                headers={"Authorization": f"Bearer {os.getenv('GROQ_API_KEY')}"},
                json={"model": GROQ_TTS_MODEL, "input": text, "voice": GROQ_TTS_VOICE, "response_format": "wav"},
                timeout=30,
            )
            if response.status_code == 200 and response.headers.get("content-type", "").startswith("audio/"):
                return response.content
            # A real API error (e.g. the model_terms_required block noted
            # above) - retrying won't change the response, fall through now.
            return None
        except (requests.exceptions.ConnectionError, requests.exceptions.Timeout):
            if attempt < GROQ_REQUEST_ATTEMPTS - 1:
                time.sleep(GROQ_RETRY_DELAY_SECONDS)
                continue
            return None
        except Exception:
            return None
    return None


def generate_speech(text: str, voice: str = DEFAULT_TTS_VOICE, allow_gemini_fallback: bool = False) -> str | None:
    """Groq TTS. Content-hash cached (PLAN §20: never regenerate the same
    text+voice pair) - unlike generate_visual_image's hash, this one does NOT
    mix in time.time(), because identical narration for the same lesson should
    always resolve to the same cached file rather than a fresh clip.

    Gemini TTS was removed along with the rest of the Google dependency: its
    free tier allowed only a handful of clips per day (observed: 5 of 21
    lessons before it ran out), so it could never actually narrate a
    curriculum. Groq TTS covers every lesson for free once its model terms are
    accepted in the Groq console; until then this returns None and the player
    falls back to the browser's own speechSynthesis, so AUDIO mode still
    speaks every lesson either way.

    `allow_gemini_fallback` is retained only so existing call sites keep
    working; it no longer does anything.

    Never raises - a missing key, rate limit, or terms-not-accepted all just
    mean "no server-side clip"; the frontend speaks it locally instead.
    """
    text = (text or "").strip()
    if not text:
        return None

    digest = hashlib.sha1(f"{text}:{voice}".encode()).hexdigest()[:16]
    filename = f"speech_{digest}.wav"
    AUDIO_DIR.mkdir(parents=True, exist_ok=True)
    target = AUDIO_DIR / filename
    if target.exists():
        return f"/static/audio/{filename}"

    audio_bytes = _generate_speech_via_groq(text)
    if not audio_bytes:
        return None

    try:
        target.write_bytes(audio_bytes)
    except OSError:
        return None

    return f"/static/audio/{filename}"
